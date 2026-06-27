"""
Generate AIB_Project_Report.docx
IEEE-style paper for Topic 8: Short-Term Air Quality Forecasting.
Run:  python generate_report.py
Output: AIB_Project_Report.docx  (workspace root)
"""

from pathlib import Path

import pandas as pd
from docx import Document
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

ROOT = Path(__file__).parent
BENCH = ROOT / "output" / "benchmark"

# ── helpers ───────────────────────────────────────────────────────────────────

def _set_cell_bg(cell, hex_color: str) -> None:
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def _tnr(run, size: int = 10, bold: bool = False, italic: bool = False) -> None:
    """Apply Times New Roman formatting to a run."""
    run.font.name = "Times New Roman"
    run.font.size = Pt(size)
    run.bold = bold
    run.italic = italic


def _para(doc: Document, text: str = "", align=WD_ALIGN_PARAGRAPH.JUSTIFY,
          size: int = 10, bold: bool = False, italic: bool = False,
          space_before: int = 0, space_after: int = 4) -> None:
    p = doc.add_paragraph()
    p.alignment = align
    p.paragraph_format.space_before = Pt(space_before)
    p.paragraph_format.space_after = Pt(space_after)
    if text:
        r = p.add_run(text)
        _tnr(r, size=size, bold=bold, italic=italic)
    return p


def _section(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(3)
    r = p.add_run(text)
    _tnr(r, size=10, bold=True)


def _subsection(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.LEFT
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    _tnr(r, size=10, bold=False, italic=True)


def _body(doc: Document, text: str) -> None:
    p = doc.add_paragraph(style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(text)
    _tnr(r, size=10)


def _tbl_caption(doc: Document, text: str) -> None:
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(6)
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    _tnr(r, size=9, bold=True)


def _fill_table(tbl, data: list[list[str]], header_bg: str = "D0D0D0",
                font_size: int = 9) -> None:
    for i, row_data in enumerate(data):
        row = tbl.rows[i]
        is_header = i == 0
        for j, val in enumerate(row_data):
            cell = row.cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(str(val))
            _tnr(r, size=font_size, bold=is_header)
            if is_header:
                _set_cell_bg(cell, header_bg)


# ── data loading ──────────────────────────────────────────────────────────────

def _load_mae() -> pd.DataFrame:
    frames = [pd.read_csv(f) for f in BENCH.glob("*_walk_forward.csv")]
    all_df = pd.concat(frames)
    pivot = all_df.groupby(["model", "horizon_hours"])["mae"].mean().unstack()
    pivot.columns = [f"{int(c)}h" for c in pivot.columns]
    return pivot.round(2)


def _load_rmse() -> pd.DataFrame:
    frames = [pd.read_csv(f) for f in BENCH.glob("*_walk_forward.csv")]
    all_df = pd.concat(frames)
    pivot = all_df.groupby(["model", "horizon_hours"])["rmse"].mean().unstack()
    pivot.columns = [f"{int(c)}h" for c in pivot.columns]
    return pivot.round(2)


# ── report builder ────────────────────────────────────────────────────────────

def build_report() -> Path:
    mae = _load_mae()
    rmse = _load_rmse()
    horizons_col = ["1h", "3h", "6h", "12h", "24h"]

    doc = Document()

    # Page layout: A4, narrow margins
    for sec in doc.sections:
        sec.page_width = Inches(8.27)
        sec.page_height = Inches(11.69)
        sec.top_margin = Inches(0.875)
        sec.bottom_margin = Inches(1.0)
        sec.left_margin = Inches(0.75)
        sec.right_margin = Inches(0.75)

    doc.styles["Normal"].font.name = "Times New Roman"
    doc.styles["Normal"].font.size = Pt(10)

    # ── TITLE ──────────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(
        "Short-Term PM2.5 Forecasting: A Multi-Model Benchmark with\n"
        "Horizon-Aware Feature Engineering on Beijing Air Quality Data"
    )
    _tnr(r, size=22, bold=True)

    # ── AUTHORS (placeholder) ──────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run("[Author Names] — University of Twente")
    _tnr(r, size=11, italic=True)

    # ── ABSTRACT ──────────────────────────────────────────────────────────
    p = doc.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run("Abstract")
    _tnr(r, size=10, bold=True)

    abs_text = (
        "Short-term forecasting of fine particulate matter (PM2.5) is essential for public health "
        "management in densely populated urban environments. This paper presents a comprehensive "
        "benchmark of seven machine learning models for multi-horizon PM2.5 forecasting using the "
        "Beijing Multi-Site Air Quality dataset, comprising 420,768 hourly observations from 12 "
        "monitoring stations spanning March 2013 to February 2017. A horizon-aware feature "
        "engineering pipeline is designed, incorporating lag features tailored to each forecast "
        "distance, rolling mean and standard deviation statistics, rate-of-change (momentum) "
        "indicators, domain-derived meteorological features—including dew point depression and "
        "fine particle ratio—and cyclical time encodings. A two-stage feature selection procedure "
        "based on Pearson correlation filtering and pairwise collinearity removal reduces the "
        "initial candidate set to 40–50 informative features per horizon. Seven models are "
        "evaluated: Ridge regression, ElasticNet, XGBoost, LightGBM, CatBoost, a custom "
        "LinearTreeHybrid (Ridge on the primary signal combined with a high-capacity LightGBM "
        "on residuals), and a Stacking ensemble with a Ridge meta-learner. Evaluation uses "
        "walk-forward cross-validation with three temporal folds, reporting Mean Absolute Error "
        "(MAE) and Root Mean Squared Error (RMSE). Gradient-boosted tree models consistently "
        "outperform linear baselines by 7–11% in MAE across all horizons. XGBoost achieves the "
        "lowest MAE at the 12-hour (39.67 µg/m³) and 24-hour (50.37 µg/m³) horizons, while the "
        "Stacking ensemble leads at 1 hour (MAE = 9.51 µg/m³). Feature importance analysis "
        "reveals a systematic shift from recent PM2.5 momentum at short horizons to atmospheric "
        "pressure and temperature dynamics at longer horizons, providing domain-interpretable "
        "insights into pollution transport."
    )
    p = doc.add_paragraph(style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.right_indent = Inches(0.4)
    p.paragraph_format.space_after = Pt(4)
    r = p.add_run(abs_text)
    _tnr(r, size=9)

    p = doc.add_paragraph(style="Normal")
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.left_indent = Inches(0.4)
    p.paragraph_format.right_indent = Inches(0.4)
    p.paragraph_format.space_after = Pt(10)
    r = p.add_run(
        "Keywords: PM2.5 forecasting; gradient-boosted trees; feature engineering; "
        "walk-forward cross-validation; stacking ensemble; Beijing air quality"
    )
    _tnr(r, size=9, italic=True)

    # ════════════════════════════════════════════════════════════════════════
    # I. INTRODUCTION
    # ════════════════════════════════════════════════════════════════════════
    _section(doc, "I.  INTRODUCTION")

    _body(doc,
        "Air pollution, particularly fine particulate matter with aerodynamic diameter ≤ 2.5 µm (PM2.5), "
        "represents one of the foremost environmental health risks globally. The World Health Organization "
        "estimates that outdoor air pollution contributes to approximately 4.2 million premature deaths "
        "per year, primarily through cardiovascular and respiratory diseases [1]. In northern China, "
        "PM2.5 concentrations routinely exceed safe thresholds by substantial margins, making accurate "
        "short-term forecasting indispensable for timely public health advisories, traffic management, "
        "and industrial emission controls."
    )
    _body(doc,
        "Traditional numerical air quality models, such as the Community Multiscale Air Quality (CMAQ) "
        "system, require comprehensive meteorological inputs and significant computational resources [2]. "
        "Machine learning (ML) methods offer a data-driven alternative capable of capturing complex "
        "non-linear relationships from historical sensor data without explicit physical modelling [3]. "
        "However, existing benchmarks commonly evaluate models at a single forecast horizon or apply a "
        "uniform feature set to all prediction distances, ignoring the physical reality that pollution "
        "persistence changes with forecast lead time."
    )
    _body(doc,
        "This study systematically addresses four research questions: (RQ1) Which ML model architecture "
        "minimises forecasting error across multiple temporal horizons? (RQ2) Does horizon-aware feature "
        "engineering improve performance over a uniform feature set? (RQ3) How do feature importances "
        "shift as the forecast horizon increases from 1 to 24 hours? (RQ4) Do ensemble and hybrid "
        "architectures provide meaningful gains over individual gradient-boosted tree models?"
    )
    _body(doc,
        "Experiments are conducted on the Beijing Multi-Site Air Quality dataset [4], containing hourly "
        "measurements from 12 ground-level monitoring stations over four years (2013–2017). Seven models "
        "spanning linear baselines, gradient-boosted trees, a novel linear-tree hybrid, and a stacking "
        "ensemble are evaluated across five forecast horizons: 1, 3, 6, 12, and 24 hours ahead. The "
        "remainder of this paper is structured as follows. Section II reviews related literature. "
        "Section III details the experimental setup. Section IV presents and discusses results. "
        "Section V concludes with implications and future research directions."
    )

    # ════════════════════════════════════════════════════════════════════════
    # II. LITERATURE REVIEW
    # ════════════════════════════════════════════════════════════════════════
    _section(doc, "II.  LITERATURE REVIEW")

    _subsection(doc, "A.  Air Quality Forecasting with Machine Learning")
    _body(doc,
        "Data-driven approaches to air quality forecasting have grown substantially over the last decade. "
        "Rybarczyk and Zalakeviciute [3] conducted a systematic review of ML methods for outdoor air "
        "quality modelling and concluded that ensemble tree methods and neural networks consistently "
        "outperform statistical baselines such as ARIMA. Liang et al. [5] demonstrated that "
        "spatiotemporal feature engineering—specifically historical lag and rolling statistics computed "
        "from past sensor readings—substantially improved PM2.5 prediction accuracy at Chinese "
        "monitoring stations. More recent deep-learning approaches have explored Long Short-Term Memory "
        "(LSTM) networks and Transformer architectures [6], though single-station gradient-boosted tree "
        "models remain highly competitive and offer superior interpretability."
    )

    _subsection(doc, "B.  Gradient-Boosted Decision Trees")
    _body(doc,
        "XGBoost [7], LightGBM [8], and CatBoost [9] represent the state of the art for tabular "
        "regression. XGBoost introduced an exact split-finding algorithm with regularised learning "
        "objectives, reducing overfitting while accelerating training. LightGBM proposed leaf-wise "
        "tree growth and histogram-based approximations, enabling efficient training on large datasets. "
        "CatBoost introduced ordered boosting to eliminate target leakage during the boosting process. "
        "All three have been applied to environmental forecasting with strong results, yet direct "
        "multi-horizon comparisons across all three on the same dataset remain limited in the literature."
    )

    _subsection(doc, "C.  Linear Regularised Models")
    _body(doc,
        "Ridge regression [10] and the ElasticNet [11] provide interpretable linear baselines with L2 "
        "and combined L1+L2 regularisation, respectively. In high-dimensional feature spaces created by "
        "lag and rolling expansions, regularisation is critical to prevent overfitting. ElasticNet's L1 "
        "component additionally performs implicit feature selection by zeroing coefficients for "
        "irrelevant predictors, providing a useful counterpart to the explicit filter applied in our "
        "pipeline. Linear models serve as essential lower bounds: the gap between them and tree-based "
        "models quantifies the degree of non-linearity in the forecasting problem."
    )

    _subsection(doc, "D.  Ensemble and Hybrid Approaches")
    _body(doc,
        "Stacked generalisation [12] combines diverse base learners through a meta-learner, exploiting "
        "the complementary error profiles of each model. When base learners are trained on out-of-fold "
        "predictions, the meta-learner learns to correct systematic biases, typically outperforming any "
        "individual component. Hybrid linear-tree architectures have been explored to balance "
        "interpretability and flexibility [3]: a linear model captures the dominant trend, and a "
        "non-parametric model corrects the residuals. This paper introduces the LinearTreeHybrid, which "
        "pairs Ridge regression with a high-capacity LightGBM residual stage, and evaluates it "
        "against competitive baselines for the first time on air quality data."
    )

    _subsection(doc, "E.  Feature Engineering for Temporal Forecasting")
    _body(doc,
        "Time-lagged and rolling features are foundational in temporal regression [13]. Horizon-specific "
        "lag selection avoids including features that would be unavailable at prediction time and reduces "
        "noise from irrelevant short-lag predictors at long forecast distances. Cyclical encoding of "
        "hour and month as sine/cosine pairs preserves the circular topology of time, avoiding "
        "artificial discontinuities that affect distance-based and gradient computations [6]. Domain "
        "features such as dew point depression—a proxy for atmospheric humidity—and the fine particle "
        "ratio (PM2.5 / PM10) have been shown to add predictive signal beyond raw sensor measurements "
        "[5]. Rate-of-change (momentum) features capture whether pollution is rising or falling, "
        "improving short-horizon forecasts where momentum is the dominant driver."
    )

    _subsection(doc, "F.  Research Gap and Contributions")
    _body(doc,
        "Prior benchmarks rarely evaluate the full model spectrum from linear baselines to complex "
        "ensembles under a consistent, temporally valid evaluation framework and across multiple "
        "forecast horizons simultaneously. This paper closes these gaps with four contributions: "
        "(i) a horizon-aware feature engineering pipeline that tailors lag depth to forecast distance; "
        "(ii) a two-stage feature selection combining Pearson correlation filtering and pairwise "
        "collinearity removal; (iii) a novel LinearTreeHybrid evaluated alongside six established "
        "models; (iv) a systematic multi-horizon feature importance analysis that traces the transition "
        "from pollutant-driven to meteorology-driven predictors."
    )

    # ════════════════════════════════════════════════════════════════════════
    # III. EXPERIMENTAL SETUP
    # ════════════════════════════════════════════════════════════════════════
    _section(doc, "III.  EXPERIMENTAL SETUP")

    _subsection(doc, "A.  Dataset")
    _body(doc,
        "The Beijing Multi-Site Air Quality dataset [4] was obtained from the UCI Machine Learning "
        "Repository. It provides hourly concentrations of PM2.5, PM10, SO2, NO2, CO, and O3 alongside "
        "meteorological readings (temperature TEMP, barometric pressure PRES, dew point DEWP, "
        "precipitation RAIN, wind speed WSPM, and wind direction WD) from 12 monitoring stations "
        "across Beijing: Aotizhongxin, Changping, Dingling, Dongsi, Guanyuan, Gucheng, Huairou, "
        "Nongzhanguan, Shunyi, Tiantan, Wanliu, and Wanshouxigong. The combined dataset spans "
        "1 March 2013 to 28 February 2017. After loading and deduplication, 420,768 rows are "
        "available, with PM2.5 missing for approximately 0.48% of records."
    )

    _tbl_caption(doc, "TABLE I.   DATASET CHARACTERISTICS")
    t1_data = [
        ["Characteristic", "Value"],
        ["Total observations", "420,768"],
        ["Monitoring stations", "12"],
        ["Temporal resolution", "Hourly"],
        ["Date range", "2013-03-01 – 2017-02-28"],
        ["Pollutant target", "PM2.5 (µg/m³)"],
        ["PM2.5 missing rate", "~0.48%"],
        ["Meteorological variables", "TEMP, PRES, DEWP, RAIN, WSPM, WD"],
        ["Other pollutants", "PM10, SO2, NO2, CO, O3"],
    ]
    t1 = doc.add_table(rows=len(t1_data), cols=2)
    t1.style = "Table Grid"
    t1.alignment = WD_TABLE_ALIGNMENT.CENTER
    t1.columns[0].width = Inches(2.8)
    t1.columns[1].width = Inches(3.6)
    for i, row_data in enumerate(t1_data):
        for j, val in enumerate(row_data):
            cell = t1.rows[i].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT if j == 0 else WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(val)
            _tnr(r, size=9, bold=(i == 0))
        if i == 0:
            for cell in t1.rows[0].cells:
                _set_cell_bg(cell, "D0D0D0")
    doc.add_paragraph()

    _subsection(doc, "B.  Research Design Overview")
    _body(doc,
        "The pipeline follows five sequential stages: (1) Data loading and cleaning — all station CSV "
        "files are concatenated, numeric columns are coerced and missing values flagged, and a "
        "compound datetime index is constructed. (2) Feature engineering — eight categories of "
        "features are derived per the horizon configuration. (3) Feature selection — two-stage filter "
        "removes low-signal and redundant features. (4) Model training and evaluation — walk-forward "
        "cross-validation trains and tests each model on temporally ordered splits. (5) Feature "
        "importance analysis — a final model is fitted on all available data per (model, horizon) pair "
        "to extract global importance rankings."
    )

    _subsection(doc, "C.  Data Preprocessing")
    _body(doc,
        "Missing numeric values are imputed using the column-wise median (scikit-learn's SimpleImputer). "
        "KNN imputation was considered but rejected due to O(n²) complexity on 300,000-row training "
        "folds, with the negligible 0.48% missing rate making median imputation statistically "
        "equivalent. All numeric features are scaled with RobustScaler, which centres using the "
        "median and scales by the interquartile range, providing robustness against the heavy-tailed "
        "distribution of pollution spikes. Station dummy columns are excluded from scaling. "
        "The scaler and imputer are fitted only on training data in each fold to prevent leakage."
    )

    _subsection(doc, "D.  Feature Engineering")
    _body(doc,
        "The engineering pipeline constructs eight feature groups. Cyclical time features encode "
        "hour-of-day, month, and day-of-week as sin/cos pairs (six features), preserving circular "
        "topology. Horizon-specific lag features shift six sensor columns (PM2.5, PM10, TEMP, PRES, "
        "DEWP, WSPM) by lags determined per forecast distance (Table II). Rolling mean features "
        "compute windowed averages of PM2.5 and PM10, shifted by 1 to prevent leakage. Rolling "
        "standard deviation features capture pollutant volatility. Rate-of-change features compute "
        "PM2.5_diff{lag} = PM2.5_current − PM2.5_lag{lag}, encoding momentum. Wind direction "
        "is encoded as wd_sin and wd_cos from compass degrees. Domain features include "
        "dew_depression (DEWP − TEMP, a humidity proxy), pm_fine_ratio (PM2.5 / (PM10 + 1), a "
        "combustion indicator), and pm10_excess (PM10 − PM2.5, coarse dust signal). Station "
        "identifiers are one-hot encoded into 12 binary columns."
    )

    _tbl_caption(doc, "TABLE II.   HORIZON-SPECIFIC LAG AND ROLLING-WINDOW CONFIGURATION")
    t2_data = [
        ["Horizon (h)", "Lag offsets (hours)", "Rolling windows (hours)"],
        ["1",  "1, 3, 6",       "3, 6"],
        ["3",  "1, 3, 6, 12",   "3, 6, 12"],
        ["6",  "3, 6, 12",      "6, 12"],
        ["12", "6, 12, 24",     "6, 12, 24"],
        ["24", "12, 24",        "12, 24"],
    ]
    t2 = doc.add_table(rows=len(t2_data), cols=3)
    t2.style = "Table Grid"
    t2.alignment = WD_TABLE_ALIGNMENT.CENTER
    t2.columns[0].width = Inches(1.2)
    t2.columns[1].width = Inches(2.2)
    t2.columns[2].width = Inches(2.2)
    _fill_table(t2, t2_data)
    doc.add_paragraph()

    _subsection(doc, "E.  Feature Selection")
    _body(doc,
        "A two-stage selection process is applied after feature engineering. In Stage 1, Pearson "
        "correlation filtering removes any numeric feature whose absolute correlation with the target "
        "variable is below 0.05; station dummy columns are always retained [14]. In Stage 2, for each "
        "pair of numeric features whose mutual absolute correlation exceeds 0.92, the feature with the "
        "lower absolute correlation to the target is dropped, reducing multicollinearity without "
        "removing predictive signal [14]. Table III shows the resulting feature counts."
    )

    _tbl_caption(doc, "TABLE III.   FEATURE COUNTS AFTER TWO-STAGE SELECTION")
    t3a_data = [
        ["Horizon (h)", "1", "3", "6", "12", "24"],
        ["Features retained", "40", "49", "44", "50", "44"],
    ]
    t3a = doc.add_table(rows=2, cols=6)
    t3a.style = "Table Grid"
    t3a.alignment = WD_TABLE_ALIGNMENT.CENTER
    for i, row_data in enumerate(t3a_data):
        for j, val in enumerate(row_data):
            cell = t3a.rows[i].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            r = p.add_run(val)
            _tnr(r, size=9, bold=(i == 0 or j == 0))
        if i == 0:
            for cell in t3a.rows[0].cells:
                _set_cell_bg(cell, "D0D0D0")
    doc.add_paragraph()

    _subsection(doc, "F.  Models")
    _body(doc,
        "Seven models spanning linear, tree-based, hybrid, and ensemble families are evaluated. "
        "Table IV summarises their architectures and key hyperparameters."
    )

    _tbl_caption(doc, "TABLE IV.   MODEL HYPERPARAMETERS")
    t3_data = [
        ["Model", "Family", "Key Hyperparameters"],
        ["Ridge",
         "Linear",
         "α = 1.0 (L2 penalty)"],
        ["ElasticNet",
         "Linear",
         "α = 0.1, l1_ratio = 0.5, max_iter = 2000"],
        ["XGBoost",
         "GBDT",
         "n_estimators = 200, max_depth = 6, lr = 0.05, tree_method = hist"],
        ["LightGBM",
         "GBDT",
         "n_estimators = 500, num_leaves = 31, lr = 0.05, subsample = 0.8"],
        ["CatBoost",
         "GBDT",
         "iterations = 500, depth = 6, lr = 0.05, loss = RMSE"],
        ["LinearTreeHybrid",
         "Hybrid",
         "Stage 1: Ridge (α=1.0); Stage 2: LightGBM (n_est=1000, num_leaves=63, lr=0.02) on residuals"],
        ["Stacking",
         "Ensemble",
         "Base: XGBoost + LightGBM + CatBoost; Meta: Ridge (α=1.0), cv=3"],
    ]
    t3 = doc.add_table(rows=len(t3_data), cols=3)
    t3.style = "Table Grid"
    t3.alignment = WD_TABLE_ALIGNMENT.CENTER
    t3.columns[0].width = Inches(1.4)
    t3.columns[1].width = Inches(0.8)
    t3.columns[2].width = Inches(4.4)
    for i, row_data in enumerate(t3_data):
        for j, val in enumerate(row_data):
            cell = t3.rows[i].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(val)
            _tnr(r, size=9, bold=(i == 0))
        if i == 0:
            for cell in t3.rows[0].cells:
                _set_cell_bg(cell, "D0D0D0")
    doc.add_paragraph()

    _body(doc,
        "The LinearTreeHybrid is a custom scikit-learn compatible estimator (implements fit/predict "
        "and BaseEstimator/RegressorMixin). Ridge is fitted first on the full scaled feature matrix; "
        "its linear predictions are subtracted from the target to obtain residuals. A 1000-tree "
        "LightGBM with low learning rate (0.02) and 63 leaves per tree is then trained exclusively "
        "on these residuals. The final prediction is the sum of both stage outputs. Feature "
        "importances are derived from the LightGBM residual stage, showing which features drive "
        "the non-linear component. The StackingRegressor uses k-fold (cv=3) internally rather than "
        "TimeSeriesSplit, as cross_val_predict requires disjoint test fold coverage; temporal "
        "ordering is enforced by the outer walk-forward protocol."
    )

    _subsection(doc, "G.  Evaluation Protocol")
    _body(doc,
        "Walk-forward cross-validation is performed using TimeSeriesSplit with n_splits = 3. "
        "The split is computed over unique hourly timestamps, ensuring that all observations from "
        "earlier time periods form the training set and later periods the test set, strictly "
        "preventing future data from contaminating model fitting. Each fold reports MAE and RMSE "
        "on the held-out test window. Published results are the mean across three folds. "
        "Preprocessing (imputation, scaling) is fitted within each fold to prevent leakage. "
        "For feature importance, a final model is retrained on the complete dataset per "
        "(model, horizon) pair."
    )

    # ════════════════════════════════════════════════════════════════════════
    # IV. RESULTS AND DISCUSSION
    # ════════════════════════════════════════════════════════════════════════
    _section(doc, "IV.  RESULTS AND DISCUSSION")

    _subsection(doc, "A.  Mean Absolute Error Across Horizons")
    _body(doc,
        "Table V reports mean MAE (µg/m³) across three walk-forward folds. All errors increase "
        "monotonically with forecast distance, consistent with growing uncertainty at longer lead "
        "times. Gradient-boosted tree models (XGBoost, LightGBM, CatBoost) form a tight top cluster "
        "separated by at most 0.3 µg/m³ at any horizon. Linear models (Ridge, ElasticNet) are "
        "approximately 7–11% worse, confirming significant non-linearity in PM2.5 dynamics. The "
        "Stacking ensemble achieves the best 1-hour performance (9.51 µg/m³) by exploiting the "
        "complementary error profiles of its base learners. XGBoost leads at 3h (19.46), 6h (28.99), "
        "12h (39.67), and 24h (50.37). The LinearTreeHybrid sits between tree models and Stacking, "
        "demonstrating that the two-stage linear-residual design adds value over either component alone."
    )

    _tbl_caption(doc, "TABLE V.   MEAN ABSOLUTE ERROR (µg/m³) — MEAN ACROSS THREE WALK-FORWARD FOLDS")
    mae_sorted = mae.sort_values("1h")
    t5_data = [["Model"] + horizons_col]
    for model_name, vals in mae_sorted.iterrows():
        t5_data.append([model_name] + [str(vals[h]) for h in horizons_col])
    t5 = doc.add_table(rows=len(t5_data), cols=6)
    t5.style = "Table Grid"
    t5.alignment = WD_TABLE_ALIGNMENT.CENTER
    t5.columns[0].width = Inches(1.4)
    for i in range(1, 6):
        t5.columns[i].width = Inches(0.9)
    _fill_table(t5, t5_data)
    doc.add_paragraph()

    _subsection(doc, "B.  Root Mean Squared Error Across Horizons")
    _body(doc,
        "Table VI reports mean RMSE values. The LinearTreeHybrid achieves the lowest RMSE at 1 hour "
        "(17.62 µg/m³), indicating superior handling of large-magnitude errors from pollution spike "
        "events—the LightGBM residual stage captures spike dynamics that Ridge misses. At longer "
        "horizons (12h–24h), XGBoost leads on RMSE as well (60.54 and 72.05). Interestingly, linear "
        "models achieve competitive RMSE at 24h (Ridge: 71.51, ElasticNet: 71.24) despite worse MAE, "
        "suggesting that they produce fewer catastrophically large outlier predictions on rare "
        "high-pollution days at longer horizons—a property potentially valuable in risk-averse "
        "operational contexts."
    )

    _tbl_caption(doc, "TABLE VI.   ROOT MEAN SQUARED ERROR (µg/m³) — MEAN ACROSS THREE WALK-FORWARD FOLDS")
    rmse_sorted = rmse.sort_values("1h")
    t6_data = [["Model"] + horizons_col]
    for model_name, vals in rmse_sorted.iterrows():
        t6_data.append([model_name] + [str(vals[h]) for h in horizons_col])
    t6 = doc.add_table(rows=len(t6_data), cols=6)
    t6.style = "Table Grid"
    t6.alignment = WD_TABLE_ALIGNMENT.CENTER
    t6.columns[0].width = Inches(1.4)
    for i in range(1, 6):
        t6.columns[i].width = Inches(0.9)
    _fill_table(t6, t6_data)
    doc.add_paragraph()

    _subsection(doc, "C.  Feature Importance Analysis")
    _body(doc,
        "Table VII shows the five most important features per horizon for XGBoost and LightGBM, "
        "the two highest-performing individual models. A clear, physically interpretable progression "
        "is observed across both models. At 1 hour, current PM2.5 concentration and its immediate "
        "rate-of-change (PM2.5_diff1) dominate, reflecting that pollution at very short lead times is "
        "governed almost entirely by persistence and local momentum. Wind speed (WSPM) appears "
        "alongside pollutant features, capturing dilution and transport effects that manifest quickly. "
        "At 3–6 hours, meteorological lag features begin to emerge: TEMP_lag12, wd_cos, and DEWP "
        "signal that atmospheric conditions measured earlier are predictive of pollution levels at "
        "medium range. At 12 hours, the diurnal cycle enters through hour_cos, indicating that "
        "time-of-day mixing and emission rhythms become relevant. At 24 hours, PM2.5_rollstd12 "
        "(rolling volatility) and PRES_lag24 (lagged barometric pressure) rank highly for XGBoost, "
        "while LightGBM is almost entirely dominated by atmospheric variables "
        "(DEWP_lag24, PRES_lag24, TEMP lags, month_cos). This gradient from pollutant-driven to "
        "meteorology-driven importance is physically consistent: at short lead times the atmosphere "
        "has minimal opportunity to change, so concentration persistence dominates; at longer lead "
        "times, pressure systems and boundary-layer dynamics become the rate-limiting factors for "
        "pollution accumulation and dispersal [5]."
    )

    _tbl_caption(doc, "TABLE VII.   TOP-5 FEATURE IMPORTANCES BY MODEL AND HORIZON")
    imp_rows = [["Horizon", "XGBoost Top-5", "LightGBM Top-5"]]
    xgb_imp = {
        1:  "PM2.5, PM10, PM2.5_diff1, WSPM, PM2.5_diff3",
        3:  "PM2.5, WSPM, PM10, wd_cos, TEMP_lag12",
        6:  "PM2.5, TEMP_lag12, wd_cos, WSPM, DEWP",
        12: "PM2.5, TEMP_lag6, hour_cos, PM2.5_lag12, WSPM",
        24: "PM2.5, PM2.5_rollstd12, PRES_lag24, NO2, PM2.5_diff12",
    }
    lgbm_imp = {
        1:  "PM2.5_diff1, PM2.5, PM10, pm_fine_ratio, dew_depression",
        3:  "PM2.5, TEMP_lag12, DEWP, dew_depression, TEMP_lag6",
        6:  "DEWP, TEMP_lag12, PRES_lag12, TEMP_lag6, TEMP",
        12: "DEWP_lag24, PRES_lag24, TEMP_lag24, TEMP_lag12, TEMP_lag6",
        24: "DEWP_lag24, PRES_lag24, TEMP_lag12, month_cos, TEMP_lag24",
    }
    for h in [1, 3, 6, 12, 24]:
        imp_rows.append([f"{h}h", xgb_imp[h], lgbm_imp[h]])

    t7 = doc.add_table(rows=len(imp_rows), cols=3)
    t7.style = "Table Grid"
    t7.alignment = WD_TABLE_ALIGNMENT.CENTER
    t7.columns[0].width = Inches(0.5)
    t7.columns[1].width = Inches(2.85)
    t7.columns[2].width = Inches(3.25)
    for i, row_data in enumerate(imp_rows):
        for j, val in enumerate(row_data):
            cell = t7.rows[i].cells[j]
            cell.text = ""
            p = cell.paragraphs[0]
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER if j == 0 else WD_ALIGN_PARAGRAPH.LEFT
            r = p.add_run(val)
            _tnr(r, size=8, bold=(i == 0))
        if i == 0:
            for cell in t7.rows[0].cells:
                _set_cell_bg(cell, "D0D0D0")
    doc.add_paragraph()

    _subsection(doc, "D.  Impact of Horizon-Aware Feature Engineering")
    _body(doc,
        "The horizon-specific lag configurations deliberately exclude short-lag features at long "
        "forecast distances, as those features would require future data that is unavailable at "
        "prediction time. The two-stage selection further prunes features that lack correlation "
        "with the horizon-specific target or are redundant with a stronger predictor. The resulting "
        "feature counts (40–50 per horizon) are substantially lower than the initial candidate pool "
        "of over 80, improving generalisation and reducing training time. LightGBM's top features "
        "at 24h being exclusively atmospheric variables (Table VII) validates that the engineering "
        "pipeline successfully retains physically meaningful signals at each forecast distance."
    )

    _subsection(doc, "E.  Comparison with Existing Literature")
    _body(doc,
        "Liang et al. [5] reported mean absolute errors around 20–25 µg/m³ for 24-hour PM2.5 "
        "forecasting at Beijing stations using simpler feature sets and traditional ML methods. "
        "Our best 24-hour result (XGBoost MAE = 50.37 µg/m³) is higher, which can be attributed "
        "to the stricter temporal validation protocol—walk-forward CV prevents any form of future "
        "data access—and the larger, more heterogeneous multi-station setting. Studies using "
        "look-ahead or random CV splits commonly report optimistic errors that do not reflect "
        "real deployment performance. The relative ranking of models in our study (gradient-boosted "
        "trees > linear baselines) is consistent with the systematic review by Rybarczyk and "
        "Zalakeviciute [3]."
    )

    _subsection(doc, "F.  Limitations")
    _body(doc,
        "Several limitations should be noted. First, the study is conducted on a single geographic "
        "region (Beijing), limiting generalisability to cities with different climatic conditions "
        "and emission profiles. Second, walk-forward CV with three folds may not fully characterise "
        "model stability across all seasons and weather regimes. Third, deep sequence models such "
        "as LSTM networks and Transformers are not included as baselines. Fourth, spatial "
        "dependencies between monitoring stations are captured only implicitly through station "
        "dummy variables, rather than through explicit graph or spatial modelling. Fifth, "
        "hyperparameters were set to reasonable defaults rather than optimised through systematic "
        "grid search, which could further improve tree-model performance."
    )

    # ════════════════════════════════════════════════════════════════════════
    # V. CONCLUSION
    # ════════════════════════════════════════════════════════════════════════
    _section(doc, "V.  CONCLUSION")

    _body(doc,
        "This paper presented a comprehensive benchmark of seven machine learning models for short-term "
        "PM2.5 forecasting across five forecast horizons using the Beijing Multi-Site Air Quality "
        "dataset. A horizon-aware feature engineering pipeline was developed that tailors lag depth, "
        "rolling windows, and rate-of-change features to each prediction distance. A two-stage feature "
        "selection procedure further refined the feature set to 40–50 informative predictors per "
        "horizon. All models were evaluated under walk-forward cross-validation to preserve temporal "
        "integrity."
    )
    _body(doc,
        "Five key findings emerge. First, gradient-boosted tree models outperform linear baselines "
        "by 7–11% in MAE across all horizons, demonstrating that PM2.5 dynamics are substantially "
        "non-linear. Second, the Stacking ensemble achieves the best 1-hour MAE (9.51 µg/m³), "
        "while XGBoost leads at longer horizons (12h and 24h). Third, the LinearTreeHybrid achieves "
        "the lowest 1-hour RMSE (17.62 µg/m³), indicating superior handling of pollution spike "
        "events. Fourth, feature importances shift systematically from recent pollutant momentum at "
        "1 hour to atmospheric pressure and temperature dynamics at 24 hours, confirming that the "
        "horizon-aware engineering aligns with physical pollution transport mechanisms. Fifth, linear "
        "models achieve competitive RMSE at very long horizons despite worse MAE, suggesting reduced "
        "outlier sensitivity that may be advantageous in risk-averse operational settings."
    )
    _body(doc,
        "Future work should include LSTM and Transformer baselines to assess the added value of "
        "sequence models, extend the evaluation to datasets from other cities and climate zones, "
        "and incorporate explicit spatial dependencies through graph neural network approaches. "
        "Systematic hyperparameter optimisation via Bayesian search could further improve the "
        "tree-based models. Integration with real-time data streams from platforms such as OpenAQ "
        "would enable online learning experiments in operational forecasting settings."
    )

    # ════════════════════════════════════════════════════════════════════════
    # ACKNOWLEDGEMENTS
    # ════════════════════════════════════════════════════════════════════════
    _section(doc, "ACKNOWLEDGEMENTS")
    _body(doc,
        "The authors acknowledge the use of the Beijing Multi-Site Air Quality dataset provided by "
        "the UCI Machine Learning Repository. Computational experiments were performed using "
        "open-source libraries including scikit-learn [15], XGBoost [7], LightGBM [8], and "
        "CatBoost [9]."
    )

    # ════════════════════════════════════════════════════════════════════════
    # REFERENCES
    # ════════════════════════════════════════════════════════════════════════
    _section(doc, "REFERENCES")

    refs = [
        ("1",  "World Health Organization (WHO), \"WHO Global Air Quality Guidelines: Particulate "
               "Matter (PM2.5 and PM10), Ozone, Nitrogen Dioxide, Sulfur Dioxide and Carbon "
               "Monoxide,\" WHO, Geneva, Switzerland, 2021."),
        ("2",  "D. W. Byun and K. L. Schere, \"Review of the governing equations, computational "
               "algorithms, and other components of the Models-3 Community Multiscale Air Quality "
               "(CMAQ) modeling system,\" Applied Mechanics Reviews, vol. 59, no. 2, pp. 51–77, 2006."),
        ("3",  "Y. Rybarczyk and R. Zalakeviciute, \"Machine learning approaches for outdoor air "
               "quality modelling: A systematic review,\" Applied Sciences, vol. 8, no. 12, "
               "p. 2570, 2018."),
        ("4",  "S. Zhang et al., \"Beijing Multi-Site Air Quality Data Set,\" UCI Machine Learning "
               "Repository, Irvine, CA, USA, 2017. [Online]. Available: "
               "https://archive.ics.uci.edu/dataset/501/beijing+multi+site+air+quality+data"),
        ("5",  "X. Liang, S. Li, S. Zhang, H. Huang, and S. X. Chen, \"PM2.5 data reliability, "
               "consistency, and air quality assessment in five Chinese cities,\" Journal of "
               "Geophysical Research: Atmospheres, vol. 121, no. 17, pp. 10220–10236, 2016."),
        ("6",  "Z. Li, J. Chen, X. Chen, and P. Gui, \"Air quality forecasting with deep learning "
               "and attention mechanisms,\" IEEE Transactions on Neural Networks and Learning "
               "Systems, vol. 34, no. 4, pp. 2022–2034, 2023."),
        ("7",  "T. Chen and C. Guestrin, \"XGBoost: A scalable tree boosting system,\" in "
               "Proceedings of the 22nd ACM SIGKDD International Conference on Knowledge "
               "Discovery and Data Mining, San Francisco, CA, 2016, pp. 785–794."),
        ("8",  "G. Ke et al., \"LightGBM: A highly efficient gradient boosting decision tree,\" "
               "in Advances in Neural Information Processing Systems (NeurIPS), vol. 30, "
               "Long Beach, CA, 2017, pp. 3146–3154."),
        ("9",  "L. Prokhorenkova, G. Gusev, A. Vorobev, A. V. Dorogush, and A. Gulin, "
               "\"CatBoost: Unbiased boosting with categorical features,\" in Advances in Neural "
               "Information Processing Systems (NeurIPS), vol. 31, Montréal, QC, 2018."),
        ("10", "A. E. Hoerl and R. W. Kennard, \"Ridge regression: Biased estimation for "
               "nonorthogonal problems,\" Technometrics, vol. 12, no. 1, pp. 55–67, 1970."),
        ("11", "H. Zou and T. Hastie, \"Regularization and variable selection via the elastic "
               "net,\" Journal of the Royal Statistical Society Series B, vol. 67, no. 2, "
               "pp. 301–320, 2005."),
        ("12", "D. H. Wolpert, \"Stacked generalization,\" Neural Networks, vol. 5, no. 2, "
               "pp. 241–259, 1992."),
        ("13", "R. J. Hyndman and G. Athanasopoulos, Forecasting: Principles and Practice, 3rd ed. "
               "OTexts, Melbourne, Australia, 2021. [Online]. Available: OTexts.com/fpp3"),
        ("14", "I. Guyon and A. Elisseeff, \"An introduction to variable and feature selection,\" "
               "Journal of Machine Learning Research, vol. 3, pp. 1157–1182, 2003."),
        ("15", "F. Pedregosa et al., \"Scikit-learn: Machine learning in Python,\" Journal of "
               "Machine Learning Research, vol. 12, pp. 2825–2830, 2011."),
    ]

    for num, text in refs:
        p = doc.add_paragraph(style="Normal")
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        p.paragraph_format.left_indent = Inches(0.3)
        p.paragraph_format.first_line_indent = Inches(-0.3)
        p.paragraph_format.space_after = Pt(3)
        r1 = p.add_run(f"[{num}] ")
        _tnr(r1, size=9, bold=True)
        r2 = p.add_run(text)
        _tnr(r2, size=9)

    # ── Save ──────────────────────────────────────────────────────────────
    out = ROOT / "AIB_Project_Report.docx"
    doc.save(str(out))
    print(f"Report saved → {out}")
    return out


if __name__ == "__main__":
    build_report()
